import streamlit as st
import matplotlib.pyplot as plt
from collections import Counter
from io import BytesIO
from crew_orchestrator import run_check
from pathway_monitor import check_for_updates
from docx import Document

st.set_page_config(page_title="Smart Doc Checker Agent", layout="wide")

st.title("📄 Smart Doc Checker Agent")
st.markdown("Upload multiple documents to find contradictions, overlaps, and ambiguities.")

# ---------- Report Generators ----------
def create_txt_report(result):
    """Generate a human-readable .txt report."""
    lines = []
    lines.append("Smart Doc Checker Report")
    lines.append("=" * 30 + "\n")

    conflicts = result.get("conflicts", [])
    suggestions = result.get("suggestions", [])

    lines.append("CONTRADICTIONS:\n")
    if conflicts:
        for i, c in enumerate(conflicts, 1):
            lines.append(f"{i}. [{c['type'].capitalize()}] {c['description']}")
    else:
        if suggestions:
            for i, s in enumerate(suggestions[:2], 1):
                lines.append(f"{i}. [Potential] {s}")
        else:
            lines.append("No contradictions found.")

    lines.append("\nSUGGESTIONS:\n")
    if suggestions:
        for s in suggestions:
            lines.append(f"- {s}")
    else:
        lines.append("No suggestions available.")

    return "\n".join(lines)


def create_docx_report(result):
    """Generate a .docx report with formatting."""
    doc = Document()
    doc.add_heading("Smart Doc Checker Report", level=1)

    conflicts = result.get("conflicts", [])
    suggestions = result.get("suggestions", [])

    # Contradictions
    doc.add_heading("Contradictions", level=2)
    if conflicts:
        for i, c in enumerate(conflicts, 1):
            doc.add_paragraph(f"{i}. [{c['type'].capitalize()}] {c['description']}")
    else:
        if suggestions:
            for i, s in enumerate(suggestions[:2], 1):
                doc.add_paragraph(f"{i}. [Potential] {s}")
        else:
            doc.add_paragraph("No contradictions found.")

    # Suggestions
    doc.add_heading("Suggestions", level=2)
    if suggestions:
        for s in suggestions:
            doc.add_paragraph(f"• {s}")
    else:
        doc.add_paragraph("No suggestions available.")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ---------- File Upload ----------
uploaded_files = st.file_uploader("Upload documents", type=["txt", "md"], accept_multiple_files=True)

if uploaded_files:
    documents = [{"filename": f.name, "text": f.read().decode("utf-8")} for f in uploaded_files]

    if st.button("Run Check"):
        with st.spinner("🔎 Analyzing documents..."):
            result = run_check(documents)

        st.success("✅ Analysis complete")

        # --- Contradictions ---
        st.subheader("📊 Contradictions Found")
        conflicts = result.get("conflicts", [])
        suggestions = result.get("suggestions", [])

        if conflicts:
            for c in conflicts:
                st.markdown(f"- **{c['type'].capitalize()}**: {c['description']}")
        else:
            if suggestions:
                st.warning("No direct contradictions detected. Showing possible issues instead:")
                for s in suggestions[:2]:
                    st.markdown(f"- ⚠️ {s}")
            else:
                st.info("No contradictions found.")

        # --- Suggestions ---
        st.subheader("💡 Suggestions")
        if suggestions:
            for s in suggestions:
                st.markdown(f"✅ {s}")
        else:
            st.info("No suggestions available.")

        # --- Conflict Distribution Graph ---
        st.subheader("📈 Conflict Distribution by Type")
        conflict_types = [c["type"].capitalize() for c in conflicts]
        if conflict_types:
            counts = Counter(conflict_types)
            fig, ax = plt.subplots(figsize=(8, 5))
            ax.barh(list(counts.keys()), list(counts.values()), color="teal")
            ax.set_title("Conflict Distribution", fontsize=14)
            ax.set_xlabel("Count")
            st.pyplot(fig)
        else:
            st.info("No conflicts to display in chart.")

        # --- Download Reports ---
        st.subheader("📥 Download Report")

        # TXT download
        txt_report = create_txt_report(result)
        st.download_button(
            label="Download .txt Report",
            data=txt_report,
            file_name="doc_analysis_report.txt",
            mime="text/plain"
        )

        # DOCX download
        docx_report = create_docx_report(result)
        st.download_button(
            label="Download .docx Report",
            data=docx_report,
            file_name="doc_analysis_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# ---------- External Policy Monitor ----------
st.divider()
st.subheader("🌐 External Policy Monitor")

if st.button("Check for Updates"):
    update = check_for_updates()
    if update:
        st.warning(f"Update detected: {update}")
    else:
        st.info("No updates detected.")
